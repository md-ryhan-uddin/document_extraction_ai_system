"""
REST API Views for document processing.
"""
from django.shortcuts import render, get_object_or_404
from django.db.models import Q
from django.http import HttpResponse, FileResponse
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
import threading
import logging
import json
import csv
from io import StringIO, BytesIO

from .models import Document, Page, ContentBlock, ExtractionLog
from .serializers import (
    DocumentSerializer, DocumentListSerializer, DocumentUploadSerializer,
    PageSerializer, PageListSerializer, ContentBlockSerializer,
    ExtractionLogSerializer, SearchResultSerializer
)
from .services import DocumentProcessor

logger = logging.getLogger(__name__)


class DocumentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for documents with upload, processing, and retrieval.
    """
    queryset = Document.objects.all()
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def get_serializer_class(self):
        """Return appropriate serializer based on action"""
        if self.action == 'list':
            return DocumentListSerializer
        elif self.action == 'create':
            return DocumentUploadSerializer
        return DocumentSerializer

    def list(self, request, *args, **kwargs):
        """List all documents with processing status info"""
        response = super().list(request, *args, **kwargs)

        # Show which documents are currently processing
        processing_docs = Document.objects.filter(status='processing')
        if processing_docs.exists():
            doc_info = ', '.join([f"#{doc.id}({doc.title})" for doc in processing_docs])
            print(f"⏳ Currently processing: {doc_info}")

        return response

    def create(self, request, *args, **kwargs):
        """Upload and process a document"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        document = serializer.save()

        print(f"\n{'='*80}")
        print(f"📤 UPLOAD RECEIVED: Document #{document.id} - '{document.title}'")
        print(f"   File: {document.original_filename}")
        print(f"   Size: {document.file.size / 1024:.2f} KB")
        print(f"   Type: {document.file_type}")
        print(f"   Starting background processing...")
        print(f"{'='*80}\n")

        # Start processing in background
        processor = DocumentProcessor()
        thread = threading.Thread(
            target=processor.process_document,
            args=(document,)
        )
        thread.daemon = True
        thread.start()

        # Return document info
        response_serializer = DocumentListSerializer(document)
        return Response(
            response_serializer.data,
            status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=['post'])
    def reprocess(self, request, pk=None):
        """Reprocess a document"""
        document = self.get_object()

        print(f"\n{'='*80}")
        print(f"🔄 REPROCESS REQUEST: Document #{document.id} - '{document.title}'")
        print(f"   Previous status: {document.status}")
        print(f"{'='*80}\n")

        # Delete existing pages and content
        document.pages.all().delete()

        # Reset document status
        document.status = 'uploaded'
        document.error_message = None
        document.save()

        # Start processing
        processor = DocumentProcessor()
        thread = threading.Thread(
            target=processor.process_document,
            args=(document,)
        )
        thread.daemon = True
        thread.start()

        return Response({
            'status': 'processing',
            'message': 'Document reprocessing started'
        })

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """Cancel document processing"""
        from .services import cancellation_manager

        document = self.get_object()

        print(f"\n{'='*60}")
        print(f"CANCEL REQUEST RECEIVED for document {document.id}")
        print(f"Current status: {document.status}")
        print(f"{'='*60}\n")

        # Only allow cancellation of processing documents
        if document.status != 'processing':
            print(f"ERROR: Cannot cancel document with status: {document.status}\n")
            return Response({
                'error': f'Cannot cancel document with status: {document.status}',
                'message': 'Only processing documents can be cancelled'
            }, status=status.HTTP_400_BAD_REQUEST)

        # Request cancellation
        print(f"Requesting cancellation for document {document.id}...")
        cancellation_manager.request_cancellation(document.id)

        # Verify it was set
        is_set = cancellation_manager.is_cancelled(document.id)
        print(f"Cancellation flag set: {is_set}")
        print(f"Cancelled docs set: {cancellation_manager._cancelled_docs}\n")

        return Response({
            'status': 'cancelling',
            'message': 'Cancellation requested. Processing will stop at next checkpoint.'
        })

    @action(detail=True, methods=['get'])
    def pages(self, request, pk=None):
        """Get all pages for a document"""
        document = self.get_object()
        pages = document.pages.all()

        # Check if we want detailed content
        include_content = request.query_params.get('include_content', 'false').lower() == 'true'

        if include_content:
            serializer = PageSerializer(pages, many=True)
        else:
            serializer = PageListSerializer(pages, many=True)

        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Download extracted content in various formats"""
        from django.http import HttpResponse
        import json
        import csv
        from io import StringIO

        document = self.get_object()
        format_type = request.query_params.get('format', 'txt')

        # Get all pages and content
        pages = document.pages.all().prefetch_related('content_blocks')

        if format_type == 'txt':
            # Plain text format
            content = f"Document: {document.title}\n"
            content += f"{'='*60}\n\n"

            for page in pages:
                content += f"Page {page.page_number}\n"
                content += f"{'-'*60}\n"
                for block in page.content_blocks.all():
                    content += f"{block.text_content}\n\n"
                content += "\n"

            response = HttpResponse(content, content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{document.title}.txt"'
            return response

        elif format_type == 'csv':
            # CSV format for tables
            output = StringIO()
            writer = csv.writer(output)
            writer.writerow(['Page', 'Block Type', 'Content'])

            for page in pages:
                for block in page.content_blocks.all():
                    writer.writerow([page.page_number, block.block_type, block.text_content])

            response = HttpResponse(output.getvalue(), content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="{document.title}.csv"'
            return response

        elif format_type == 'json':
            # JSON format
            data = {
                'document': {
                    'title': document.title,
                    'status': document.status,
                    'total_pages': document.total_pages,
                },
                'pages': []
            }

            for page in pages:
                page_data = {
                    'page_number': page.page_number,
                    'language': page.detected_language,
                    'page_type': page.page_type,
                    'content_blocks': []
                }

                for block in page.content_blocks.all():
                    block_data = {
                        'block_number': block.block_number,
                        'block_type': block.block_type,
                        'text_content': block.text_content,
                        'confidence': block.confidence,
                        'table_data': block.table_data,
                        'form_data': block.form_data
                    }
                    page_data['content_blocks'].append(block_data)

                data['pages'].append(page_data)

            response = HttpResponse(json.dumps(data, indent=2), content_type='application/json')
            response['Content-Disposition'] = f'attachment; filename="{document.title}.json"'
            return response

        elif format_type == 'docx':
            # Word document format (requires python-docx)
            try:
                from docx import Document as DocxDocument
                from docx.shared import Inches

                doc = DocxDocument()
                doc.add_heading(document.title, 0)

                for page in pages:
                    doc.add_heading(f'Page {page.page_number}', level=1)
                    for block in page.content_blocks.all():
                        if block.block_type == 'table' and block.table_data:
                            # Add table
                            doc.add_paragraph(block.text_content)
                        else:
                            doc.add_paragraph(block.text_content)

                # Save to response
                from io import BytesIO
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
                return response
            except ImportError:
                return Response({
                    'error': 'python-docx not installed',
                    'message': 'Word document export requires python-docx package'
                }, status=status.HTTP_400_BAD_REQUEST)

        return Response({'error': 'Invalid format'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['get'], url_path='download-original')
    def download_original(self, request, pk=None):
        """Download original uploaded file"""
        from django.http import FileResponse

        document = self.get_object()

        if not document.file:
            return Response({
                'error': 'File not found',
                'message': 'Original file is not available'
            }, status=status.HTTP_404_NOT_FOUND)

        response = FileResponse(document.file.open('rb'))
        response['Content-Disposition'] = f'attachment; filename="{document.original_filename}"'
        return response

    @action(detail=True, methods=['get'], url_path='download-tables')
    def download_tables(self, request, pk=None):
        """Download only tables in specified format"""
        from django.http import HttpResponse
        import json
        import csv
        from io import StringIO, BytesIO

        document = self.get_object()
        format_type = request.query_params.get('format', 'xlsx')

        # Get all content blocks that are tables
        pages = document.pages.all().prefetch_related('content_blocks')
        table_blocks = []
        for page in pages:
            for block in page.content_blocks.filter(block_type='table'):
                if block.table_data:
                    table_blocks.append({
                        'page': page.page_number,
                        'block': block.block_number,
                        'data': block.table_data
                    })

        if not table_blocks:
            return Response({
                'error': 'No tables found',
                'message': 'This document does not contain any extracted tables'
            }, status=status.HTTP_404_NOT_FOUND)

        if format_type == 'xlsx':
            # Excel format - requires openpyxl
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill, Alignment

                wb = Workbook()
                wb.remove(wb.active)  # Remove default sheet

                for idx, table_block in enumerate(table_blocks):
                    sheet_name = f"Page{table_block['page']}_Table{table_block['block']}"
                    ws = wb.create_sheet(title=sheet_name[:31])  # Excel limits sheet names to 31 chars

                    # Parse table data and write to Excel
                    table_data = table_block['data']
                    if 'cells' in table_data:
                        # Build table from cells
                        max_row = max([cell.get('row', 0) for cell in table_data['cells']], default=0) + 1
                        max_col = max([cell.get('col', 0) for cell in table_data['cells']], default=0) + 1

                        for cell in table_data['cells']:
                            row = cell.get('row', 0) + 1
                            col = cell.get('col', 0) + 1
                            value = cell.get('text', '')
                            ws.cell(row=row, column=col, value=value)

                            # Style header row
                            if row == 1:
                                ws.cell(row=row, column=col).font = Font(bold=True)
                                ws.cell(row=row, column=col).fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

                buffer = BytesIO()
                wb.save(buffer)
                buffer.seek(0)

                response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Disposition'] = f'attachment; filename="{document.title}_tables.xlsx"'
                return response
            except ImportError:
                return Response({
                    'error': 'openpyxl not installed',
                    'message': 'Excel export requires openpyxl package'
                }, status=status.HTTP_400_BAD_REQUEST)

        elif format_type == 'csv':
            # CSV format - one file with all tables
            output = StringIO()
            writer = csv.writer(output)

            for table_block in table_blocks:
                writer.writerow([f"=== Page {table_block['page']}, Table {table_block['block']} ==="])
                table_data = table_block['data']

                if 'cells' in table_data:
                    # Build table from cells
                    max_row = max([cell.get('row', 0) for cell in table_data['cells']], default=0) + 1
                    max_col = max([cell.get('col', 0) for cell in table_data['cells']], default=0) + 1

                    # Create empty grid
                    grid = [['' for _ in range(max_col)] for _ in range(max_row)]

                    # Fill grid
                    for cell in table_data['cells']:
                        row = cell.get('row', 0)
                        col = cell.get('col', 0)
                        grid[row][col] = cell.get('text', '')

                    # Write grid to CSV
                    for row in grid:
                        writer.writerow(row)

                writer.writerow([])  # Empty row between tables

            response = HttpResponse(output.getvalue(), content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="{document.title}_tables.csv"'
            return response

        elif format_type == 'pdf':
            # PDF format - requires reportlab
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4)
                elements = []
                styles = getSampleStyleSheet()

                for table_block in table_blocks:
                    # Add heading
                    heading = Paragraph(f"Page {table_block['page']}, Table {table_block['block']}", styles['Heading2'])
                    elements.append(heading)
                    elements.append(Spacer(1, 12))

                    table_data = table_block['data']
                    if 'cells' in table_data:
                        # Build table from cells
                        max_row = max([cell.get('row', 0) for cell in table_data['cells']], default=0) + 1
                        max_col = max([cell.get('col', 0) for cell in table_data['cells']], default=0) + 1

                        # Create empty grid
                        grid = [['' for _ in range(max_col)] for _ in range(max_row)]

                        # Fill grid
                        for cell in table_data['cells']:
                            row = cell.get('row', 0)
                            col = cell.get('col', 0)
                            grid[row][col] = cell.get('text', '')

                        # Create PDF table
                        pdf_table = Table(grid)
                        pdf_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 12),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        elements.append(pdf_table)
                        elements.append(Spacer(1, 20))

                doc.build(elements)
                buffer.seek(0)

                response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{document.title}_tables.pdf"'
                return response
            except ImportError:
                return Response({
                    'error': 'reportlab not installed',
                    'message': 'PDF export requires reportlab package'
                }, status=status.HTTP_400_BAD_REQUEST)

        return Response({'error': 'Invalid format'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['get'])
    def search(self, request):
        """
        Search across all documents.

        Query params:
        - q: Search query
        - document_id: Filter by document
        - page_type: Filter by page type
        - block_type: Filter by block type
        - language: Filter by language
        """
        query = request.query_params.get('q', '')
        document_id = request.query_params.get('document_id')
        page_type = request.query_params.get('page_type')
        block_type = request.query_params.get('block_type')
        language = request.query_params.get('language')

        # Base queryset
        blocks = ContentBlock.objects.select_related('page', 'page__document').all()

        # Apply filters
        if query:
            blocks = blocks.filter(text_content__icontains=query)

        if document_id:
            blocks = blocks.filter(page__document_id=document_id)

        if page_type:
            blocks = blocks.filter(page__page_type=page_type)

        if block_type:
            blocks = blocks.filter(block_type=block_type)

        if language:
            blocks = blocks.filter(page__detected_language=language)

        # Build results
        results = []
        for block in blocks[:100]:  # Limit to 100 results
            results.append({
                'document_id': block.page.document.id,
                'document_title': block.page.document.title,
                'page_id': block.page.id,
                'page_number': block.page.page_number,
                'block_id': block.id,
                'block_type': block.block_type,
                'text_content': block.text_content[:500],  # Limit text
                'confidence': block.confidence,
            })

        serializer = SearchResultSerializer(results, many=True)
        return Response({
            'count': len(results),
            'results': serializer.data
        })


class PageViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for pages"""
    queryset = Page.objects.all()

    def get_serializer_class(self):
        """Return appropriate serializer"""
        if self.action == 'list':
            return PageListSerializer
        return PageSerializer

    @action(detail=True, methods=['get'])
    def content(self, request, pk=None):
        """Get all content blocks for a page"""
        page = self.get_object()
        blocks = page.content_blocks.all()
        serializer = ContentBlockSerializer(blocks, many=True)
        return Response(serializer.data)


class ContentBlockViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for content blocks"""
    queryset = ContentBlock.objects.all()
    serializer_class = ContentBlockSerializer

    def get_queryset(self):
        """Filter by page if provided"""
        queryset = super().get_queryset()
        page_id = self.request.query_params.get('page_id')

        if page_id:
            queryset = queryset.filter(page_id=page_id)

        return queryset

    @action(detail=True, methods=['get'], url_path='export')
    def download_content(self, request, pk=None):
        """Download individual content block in various formats"""
        from django.http import HttpResponse
        import json
        import csv
        from io import StringIO, BytesIO

        block = self.get_object()
        format_type = request.query_params.get('export_format', 'txt')
        page = block.page

        # Generate filename base
        filename_base = f"{page.document.title}_Page{page.page_number}_Block{block.block_number}"

        logger.info(f"Downloading block {block.id} (type: {block.block_type}) as {format_type}")

        try:
            if format_type == 'txt':
                # Plain text format with table formatting
                content = f"Document: {page.document.title}\n"
                content += f"Page: {page.page_number}\n"
                content += f"Block Type: {block.block_type}\n"
                content += f"{'='*80}\n\n"

                # For tables, render as formatted table
                if block.block_type == 'table' and block.table_data:
                    table_data = block.table_data

                    # Get headers
                    if 'headers' in table_data and table_data['headers']:
                        headers = [h.get('text', '') for h in table_data['headers']]
                    else:
                        headers = []

                    # Get rows
                    rows_data = []
                    if 'rows' in table_data and table_data['rows']:
                        for row in table_data['rows']:
                            if 'cells' in row:
                                row_cells = [cell.get('text', '') for cell in row['cells']]
                                rows_data.append(row_cells)

                    # Calculate column widths
                    if headers or rows_data:
                        num_cols = len(headers) if headers else (len(rows_data[0]) if rows_data else 0)
                        col_widths = [15] * num_cols  # Default width

                        # Adjust based on content
                        if headers:
                            for i, h in enumerate(headers):
                                col_widths[i] = max(col_widths[i], len(str(h)) + 2)

                        for row in rows_data:
                            for i, cell in enumerate(row):
                                if i < len(col_widths):
                                    col_widths[i] = max(col_widths[i], len(str(cell)) + 2)

                        # Render table
                        separator = '+' + '+'.join(['-' * w for w in col_widths]) + '+'

                        content += separator + '\n'

                        # Render headers
                        if headers:
                            header_row = '|' + '|'.join([h.ljust(col_widths[i]) for i, h in enumerate(headers)]) + '|'
                            content += header_row + '\n'
                            content += separator + '\n'

                        # Render rows
                        for row in rows_data:
                            row_str = '|' + '|'.join([str(cell).ljust(col_widths[i]) for i, cell in enumerate(row)]) + '|'
                            content += row_str + '\n'

                        content += separator + '\n'
                    else:
                        content += block.text_content or ''
                else:
                    # For non-table blocks, just use text content
                    content += block.text_content or ''

                response = HttpResponse(content, content_type='text/plain')
                response['Content-Disposition'] = f'attachment; filename="{filename_base}.txt"'
                return response

            if format_type == 'json':
                # JSON format
                data = {
                    'document_title': page.document.title,
                    'page_number': page.page_number,
                    'block_number': block.block_number,
                    'block_type': block.block_type,
                    'text_content': block.text_content,
                    'confidence': block.confidence,
                    'table_data': block.table_data,
                    'form_data': block.form_data
                }

                response = HttpResponse(json.dumps(data, indent=2), content_type='application/json')
                response['Content-Disposition'] = f'attachment; filename="{filename_base}.json"'
                return response

            if format_type == 'xlsx':
                # Excel format for tables
                try:
                    from openpyxl import Workbook
                    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

                    wb = Workbook()
                    ws = wb.active
                    ws.title = f"Page{page.page_number}_Block{block.block_number}"[:31]

                    if block.block_type == 'table' and block.table_data:
                        table_data = block.table_data
                        current_row = 1

                        # Add headers
                        if 'headers' in table_data and table_data['headers']:
                            for col_idx, header in enumerate(table_data['headers'], start=1):
                                cell = ws.cell(row=current_row, column=col_idx, value=header.get('text', ''))
                                # Style headers
                                cell.font = Font(bold=True, color="FFFFFF")
                                cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                                cell.alignment = Alignment(horizontal='center', vertical='center')
                                cell.border = Border(
                                    left=Side(style='thin'),
                                    right=Side(style='thin'),
                                    top=Side(style='thin'),
                                    bottom=Side(style='thin')
                                )
                            current_row += 1

                        # Add rows
                        if 'rows' in table_data and table_data['rows']:
                            for row_data in table_data['rows']:
                                if 'cells' in row_data:
                                    for col_idx, cell_data in enumerate(row_data['cells'], start=1):
                                        cell = ws.cell(row=current_row, column=col_idx, value=cell_data.get('text', ''))
                                        cell.alignment = Alignment(horizontal='left', vertical='center')
                                        cell.border = Border(
                                            left=Side(style='thin'),
                                            right=Side(style='thin'),
                                            top=Side(style='thin'),
                                            bottom=Side(style='thin')
                                        )
                                    current_row += 1

                        # Auto-adjust column widths
                        for column in ws.columns:
                            max_length = 0
                            column_letter = column[0].column_letter
                            for cell in column:
                                try:
                                    if len(str(cell.value)) > max_length:
                                        max_length = len(str(cell.value))
                                except:
                                    pass
                            adjusted_width = min(max_length + 2, 50)
                            ws.column_dimensions[column_letter].width = adjusted_width
                    else:
                        # For non-table blocks, just put text content
                        if block.text_content:
                            ws.cell(row=1, column=1, value=block.text_content)

                    buffer = BytesIO()
                    wb.save(buffer)
                    buffer.seek(0)

                    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = f'attachment; filename="{filename_base}.xlsx"'
                    return response
                except ImportError:
                    return Response({
                        'error': 'openpyxl not installed',
                        'message': 'Excel export requires openpyxl package'
                    }, status=status.HTTP_400_BAD_REQUEST)

            if format_type == 'csv':
                # CSV format
                output = StringIO()
                writer = csv.writer(output)

                if block.block_type == 'table' and block.table_data:
                    table_data = block.table_data

                    # Write headers
                    if 'headers' in table_data and table_data['headers']:
                        headers = [h.get('text', '') for h in table_data['headers']]
                        writer.writerow(headers)

                    # Write rows
                    if 'rows' in table_data and table_data['rows']:
                        for row_data in table_data['rows']:
                            if 'cells' in row_data:
                                row = [cell.get('text', '') for cell in row_data['cells']]
                                writer.writerow(row)
                else:
                    # For non-table blocks, just write text content
                    if block.text_content:
                        writer.writerow([block.text_content])

                response = HttpResponse(output.getvalue(), content_type='text/csv')
                response['Content-Disposition'] = f'attachment; filename="{filename_base}.csv"'
                return response

            if format_type == 'pdf':
                # PDF format
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib import colors
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

                    buffer = BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    elements = []
                    styles = getSampleStyleSheet()

                    # Add title
                    title = Paragraph(f"<b>{page.document.title}</b>", styles['Title'])
                    elements.append(title)
                    elements.append(Spacer(1, 12))

                    # Add metadata
                    meta = Paragraph(f"Page {page.page_number} - Block {block.block_number} ({block.block_type})", styles['Normal'])
                    elements.append(meta)
                    elements.append(Spacer(1, 12))

                    if block.block_type == 'table' and block.table_data:
                        # Render table
                        table_data = block.table_data
                        pdf_data = []

                        # Add headers
                        if 'headers' in table_data and table_data['headers']:
                            headers = [h.get('text', '') for h in table_data['headers']]
                            pdf_data.append(headers)

                        # Add rows
                        if 'rows' in table_data and table_data['rows']:
                            for row_data in table_data['rows']:
                                if 'cells' in row_data:
                                    row = [cell.get('text', '') for cell in row_data['cells']]
                                    pdf_data.append(row)

                        # Create PDF table
                        if pdf_data:
                            pdf_table = Table(pdf_data)
                            # Style the table
                            table_style = [
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 11),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                                ('TOPPADDING', (0, 0), (-1, 0), 10),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ]

                            # Alternate row colors
                            for i in range(1, len(pdf_data)):
                                if i % 2 == 0:
                                    table_style.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor('#E7E6E6')))
                                else:
                                    table_style.append(('BACKGROUND', (0, i), (-1, i), colors.white))

                            pdf_table.setStyle(TableStyle(table_style))
                            elements.append(pdf_table)
                    elif block.block_type == 'form' and block.form_data:
                        # Render form fields
                        form_data = block.form_data
                        if 'fields' in form_data:
                            for field in form_data['fields']:
                                field_text = f"<b>{field.get('field_label', field.get('field_name', 'Unknown'))}:</b> {field.get('field_value', '(empty)')}"
                                elements.append(Paragraph(field_text, styles['Normal']))
                                elements.append(Spacer(1, 6))
                    else:
                        # Plain text content
                        text = Paragraph(block.text_content or '', styles['Normal'])
                        elements.append(text)

                    doc.build(elements)
                    buffer.seek(0)

                    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{filename_base}.pdf"'
                    return response
                except ImportError:
                    return Response({
                        'error': 'reportlab not installed',
                        'message': 'PDF export requires reportlab package'
                    }, status=status.HTTP_400_BAD_REQUEST)

            # Invalid format
            return Response({
                'error': 'Invalid format',
                'message': f"Format '{format_type}' is not supported"
            }, status=status.HTTP_400_BAD_REQUEST)

        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            logger.error(f"Unexpected error in block download: {str(e)}\n{error_trace}")
            return Response({
                'error': 'Internal server error',
                'message': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ExtractionLogViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for extraction logs"""
    queryset = ExtractionLog.objects.all()
    serializer_class = ExtractionLogSerializer

    def get_queryset(self):
        """Filter by document if provided"""
        queryset = super().get_queryset()
        document_id = self.request.query_params.get('document_id')

        if document_id:
            queryset = queryset.filter(document_id=document_id)

        return queryset


# Frontend views
def home(request):
    """Home page with upload interface"""
    return render(request, 'documents/home.html')


def all_documents(request):
    """All documents page"""
    # Prepare a lightweight JSON representation of all documents so the
    # template can render them immediately without relying on an API call.
    documents_qs = Document.objects.all()
    documents = []
    for d in documents_qs:
        documents.append({
            'id': d.id,
            'title': d.title,
            'file_type': d.file_type,
            'status': d.status,
            'file_size': d.file_size,
            'total_pages': d.total_pages,
            'uploaded_at': d.uploaded_at.isoformat() if d.uploaded_at else None,
        })

    return render(request, 'documents/all_documents.html', {
        'documents_json': json.dumps(documents)
    })


def document_viewer(request, document_id):
    """Document viewer page"""
    document = get_object_or_404(Document, pk=document_id)
    return render(request, 'documents/viewer.html', {'document': document})
